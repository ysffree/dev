#!/usr/bin/env python3
# coding:utf-8

# paragraph
# p1：楷体、Times New Roman 五号、不加粗、首行不缩进
# p2: 楷体、Times New Roman 五号、加粗、首行缩进
# p3：楷体、Times New Roman 五号、加粗、首行不缩进
# p4：楷体、Times New Roman 五号、不加粗、首行缩进
# p5：楷体、Times New Roman 四号、加粗、首行不缩进
# p6：楷体、Times New Roman 10号、加粗、首行不缩进，段后距0
# p7：楷体、Times New Roman 六号 加粗 红色 居中
# p8：楷体、Times New Roman 小五 加粗 红色 居中
# pic：居中

# run
# r1：楷体GB2312、Times New Roman 小三
# r2：楷体、Times New Roman 五号 下划线
# r3：楷体、Calibri (Body) 五号
# r4：楷体、Times New Roman 六号 加粗 红色
# r5：楷体、Times New Roman 小五 加粗 红色
# r6：楷体、Times New Roman 五号 加粗
# r7：楷体、Times New Roman 五号
# r8：楷体、Calibri (Body) 五号 加粗
# r9：楷体、Times New Roman 小五 加粗
# r10：楷体、Times New Roman 8
# r11：楷体、Times New Roman 小五
# strong：楷体、Times New Roman 五号 加粗 红色



from docx import Document
from .person_info import get_person_info  # 1
# from person_info import  get_person_info
from optparse import OptionParser
import re
from .chem import insert_chem  # 2
# from chem import insert_chem
# from person_info import get_person_info
from docx.shared import Inches
import os
import json
from collections import OrderedDict
import PATH

DATA_PATH = PATH.sample_path  # 4
# DATA_PATH = 'D:\work\\try\daan'
# LOCAL_PATH = 'D:\work\\try\daan\daan\linux'
LOCAL_PATH = os.path.join(PATH.home_path, 'module', 'daan')  # 3
PIC_PATH = os.path.join(LOCAL_PATH, 'config', 'gene_picture')
TEMPLE_FILE = os.path.join(LOCAL_PATH, 'config', 'temple_daan.docx')
SAVE_PATH = '/fastzone/worker/worker_report/temp_report'


TRANSVAR = OrderedDict((
    ('mutation', '点突变'),
    ('fusion', '融合变异'),
    ('cnv', '拷贝数变异'),
    ('CNV扩增', 'copy number GAIN')
))

parser = OptionParser(usage="usage:%prog [options] arg1 arg2")
parser.add_option("-s", "--sample",
                  action="store",
                  help="样本编号")
parser.add_option("-t", "--sample_type",
                  action="store",
                  default='Plasma',
                  help="样本类型")
parser.add_option("-o", "--out_path",
                  action="store",
                  default='D:\work\\try\daan',
                  help="输出路径")
parser.add_option("-m", "--mutation_path",
                  action="store",
                  default=None,
                  help="突变基因文件路径")
# parser.add_option("-v", "--version",
#                   action="store",
#                   default='Smartlung',
#                   help="报告版本")



def insert_person_info(person_info, doc):
    for pa in doc.paragraphs:
        if pa.text == '受检者：':
            pa.add_run(person_info.get("姓  名"), 'r1')
        elif pa.text == '接收日期：':
            pa.add_run(person_info.get("样本接收日期"), 'r1')
        elif pa.text == '报告日期：':
            pa.add_run(person_info.get("报告日期"), 'r1')
        # elif pa.text == '您好！':
        elif '您好！' in pa.text:
            # pa.insert_paragraph_before('尊敬的   {0}   先生/女士：'.format(person_info["姓  名"]), 'p1')
            par = pa.insert_paragraph_before('尊敬的', 'p1')
            par.add_run('   {0}   '.format(person_info.get("姓  名")), 'r2')
            par.add_run('先生/女士：', 'r7')

    person_table = doc.tables[0]
    person_table.cell(0, 0).paragraphs[0].add_run(person_info.get("姓  名"), 'r8')
    person_table.cell(0, 1).paragraphs[0].add_run(person_info.get("性  别"), 'r8')
    person_table.cell(1, 1).paragraphs[0].add_run(person_info.get("病理诊断"), 'r3')
    person_table.cell(1, 3).paragraphs[0].add_run(person_info.get("样本类型"), 'r3')

def walk_dir(path, keyword):
    for (root, subdirs, files) in os.walk(path):
        for file in files:
            if keyword in file:
                return os.path.join(root, file)

def get_qc(sample_type, sample_path):
    qc_path = os.path.join(sample_path, 'qc')
    total_file = walk_dir(qc_path, 'R1.fastq')
    cover_file = walk_dir(qc_path, 'panel.stat')
    qc_dict = {}
    print(total_file)
    with open(total_file, encoding='utf8') as f:
        for line in f:
            if 'Total Base' in line:
                Total_raw = float(line.split(':')[2])
                deal_raw = str(int(Total_raw/1000000 * 2))
                if len(deal_raw) > 3:
                    qc_dict['测序数据'] = str(round(float(deal_raw)/1000, 1)) + 'G'
                else:
                    qc_dict['测序数据'] = deal_raw + 'M'

    with open(cover_file, encoding='utf8') as f:
        for line in f:
            if 'Mean depth' in line:
                depth_raw = float(line.split(',')[1])
                qc_dict['平均覆盖深度'] = str(int(depth_raw)) + 'X'
    if sample_type == 'Plasma':
        qc_dict['变异频率'] = '>0.1%'
    else:
        qc_dict['变异频率'] = '>1%'
    return qc_dict

def insert_qc(sample_id, sample_type, sample_path, doc):
    qc_table = doc.tables[0]
    qc_dict = get_qc(sample_type, sample_path)
    qc_table.cell(2, 2).paragraphs[0].add_run(sample_id, 'r3')
    qc_table.cell(4, 2).paragraphs[0].add_run(qc_dict.get('变异频率'), 'r3')
    qc_table.cell(5, 0).paragraphs[0].add_run(qc_dict.get('测序数据'), 'r8')
    qc_table.cell(5, 2).paragraphs[0].add_run(qc_dict.get('平均覆盖深度'), 'r3')

def load_json(json_file):
    with open(json_file, encoding='utf8') as f:
        json_data = json.load(f)
    return json_data

def insert_result(json_data, doc):
    tbl = doc.tables[1]
    gene = []
    for mutation in json_data:
        if mutation.get('gene_name') != tbl.cell(-1, 2).text or mutation.get('gene_C') != tbl.cell(-1, 6).text:
            tbl.add_row()
            tbl.cell(-1, 2).paragraphs[0].add_run(mutation.get('gene_name'), 'r4')
            tbl.cell(-1, 3).paragraphs[0].add_run(mutation.get('gene_exon'), 'r4')
            if '_' in mutation['gene_C']:
                tbl.cell(-1, 4).paragraphs[0].add_run('插入缺失变异', 'r4')
            else:
                tbl.cell(-1, 4).paragraphs[0].add_run(TRANSVAR[mutation.get('gene_mut_type')], 'r4')
            tbl.cell(-1, 5).paragraphs[0].add_run(mutation.get('gene_G'), 'r4')
            tbl.cell(-1, 6).paragraphs[0].add_run(mutation.get('gene_C'), 'r4')
            if 'CNV' in mutation['gene_P']:
                tbl.cell(-1, 7).paragraphs[0].add_run(TRANSVAR[mutation.get('gene_P')], 'r4')
            else:
                tbl.cell(-1, 7).paragraphs[0].add_run(mutation.get('gene_P'), 'r4')
            tbl.cell(-1, 8).paragraphs[0].add_run(mutation.get('gene_name'), 'r4')
            tbl.cell(-1, 9).paragraphs[0].add_run(mutation.get('gene_name'), 'r4')
            tbl.cell(-1, 10).paragraphs[0].add_run(str(mutation.get('gene_frequency')), 'r4')
# 外显子	变异类型	变异位置	cHGVS	pHGVS	覆盖深度	变异覆盖深度	突变频率

def insert_drug(json_data, doc):
    lung_drug_table = doc.tables[2]
    other_drug_table = doc.tables[3]
    drug_use_table = doc.tables[4]
    lung_drug = []
    for drug in lung_drug_table.column_cells(0):
        lung_drug.append(drug.text)
    lung_drug = ''.join(lung_drug)
   
    for drug_use in json_data:
        if 'drug_ch_name' not in drug_use:
            continue
        if '肺癌' in str(drug_use.get('drug_ch_cancer')) and str(drug_use.get('drug_ch_name')) in lung_drug:
            lung_drug_table.add_row()
            lung_drug_table.cell(-1, 1).paragraphs[0].add_run(drug_use.get('drug_ch_name'), 'r5')
            lung_drug_table.cell(-1, 2).paragraphs[0].add_run(drug_use.get('gene_name'), 'r5')
            if 'CNV' in drug_use['gene_P']:
                lung_drug_table.cell(-1, 3).paragraphs[0].add_run(TRANSVAR[drug_use.get('gene_P')], 'r5')
            else:
                lung_drug_table.cell(-1, 3).paragraphs[0].add_run('{0} {1}'.format(drug_use.get('gene_exon'), drug_use.get('gene_P')), 'r5')
            if '_' in drug_use['gene_C']:
                lung_drug_table.cell(-1, 4).paragraphs[0].add_run('插入缺失变异', 'r5')
            else:
                lung_drug_table.cell(-1, 4).paragraphs[0].add_run(TRANSVAR[drug_use.get('gene_mut_type')], 'r5')
            lung_drug_table.cell(-1, 5).paragraphs[0].add_run(drug_use.get('drug_Tips'), 'r5')

        else:
            other_drug_table.add_row()
            other_drug_table.cell(-1, 0).paragraphs[0].add_run(drug_use.get('drug_ch_name'), 'r9')
            other_drug_table.cell(-1, 1).paragraphs[0].add_run('临床研究\n{0}\n{1}'.format(drug_use.get('drug_ch_cancer'), drug_use.get('drug_summary_id')), 'r10')
            other_drug_table.cell(-1, 2).paragraphs[0].add_run(drug_use.get('gene_name'), 'r11')
            if 'CNV' in drug_use['gene_P']:
                other_drug_table.cell(-1, 3).paragraphs[0].add_run(TRANSVAR[drug_use.get('gene_P')], 'r5')
            else:
                other_drug_table.cell(-1, 3).paragraphs[0].add_run('{0} {1}'.format(drug_use.get('gene_exon'), drug_use.get('gene_P')), 'r5')
            if '_' in drug_use['gene_C']:
                other_drug_table.cell(-1, 4).paragraphs[0].add_run('插入缺失变异', 'r5')
            else:
                other_drug_table.cell(-1, 4).paragraphs[0].add_run(TRANSVAR[drug_use.get('gene_mut_type')], 'r5')
            other_drug_table.cell(-1, 5).paragraphs[0].add_run(drug_use.get('drug_Tips'), 'r5')
            drug_use_table.add_row()
            drug_use_table.cell(-1, 0).paragraphs[0].add_run(drug_use.get('drug_ch_name'), 'r6')
            drug_use_table.cell(-1, 2).paragraphs[0].add_run(drug_use.get('drug_Tips'), 'r6')
            drug_use_table.cell(-1, 4).paragraphs[0].add_run('临床研究\n{0}\n{1}'.format(drug_use.get('drug_ch_cancer'), drug_use.get('drug_summary_id')), 'r10')

def insert_drug_instruction(json_data, sample_path, doc):
    gene_file = '/mnt/11d1/worker/ctDNA_result_link/S170003167-Plasma/report-results/S170003167.somaticlist_somaticli/result_gene.json'
    #gene_file = walk_dir(sample_path, 'result_gene.json')
    print(sample_path)
    gene_json = load_json(gene_file)
    gene_list = []
    for drug in json_data:
        if drug.get('gene_name') not in gene_list:
            for gene in gene_json:
                if gene.get("基因名称") == drug.get('gene_name'):
                    gene_list.append(drug['gene_name'])
                    doc.add_paragraph()
                    doc.add_paragraph('\u25A1\t{0} {1} {2}基因突变'.format(drug.get('gene_name'), drug.get('gene_exon'), drug.get('gene_P')), 'p3')
                    doc.add_paragraph('基因变异解析', 'p3')
                    doc.add_paragraph(gene.get("基因说明"), 'p4')
                    doc.add_paragraph(style='pic').add_run().add_picture(os.path.join(PIC_PATH, '{}.png'.format(gene.get("基因名称"))), width=Inches(3.8))
                    doc.add_paragraph('基因变异与药物敏感性解析：', 'p3')
                    break
        if 'drug_ch_name' in drug:
            doc.add_paragraph('{0}：'.format(drug.get('drug_ch_name')), 'p2').add_run(drug.get('drug_summary'), 'r6')
        else:
            doc.add_paragraph('本次检测到的{0} {1} {2} {3}突变暂无临床数据提示与靶向药物有密切关系。'.format(drug.get('gene_name'), drug.get('gene_exon'), drug.get('gene_C'), drug.get('gene_P')), 'p2')
    doc.add_page_break()

def save_doc(sample_id, person_info, doc, save_path):
    if save_path:
        save_file = os.path.join(save_path, '{0}-{1}-上海达安医学检验所-SmartLung Plus-14基因.docx').format(sample_id, person_info.get("姓  名"))
        doc.save(save_file)
    else:
        save_file = os.path.join(SAVE_PATH, '{0}-{1}-上海达安医学检验所-SmartLung Plus-14基因.docx').format(sample_id, person_info.get("姓  名"))
    doc.save(save_file)
    print(save_file)


def daan_run(sample_id, sample_type, json_data, save_path=None):
    doc = Document(TEMPLE_FILE)
    sample_path = os.path.join(DATA_PATH, sample_id + '-' + sample_type)
    person_info = get_person_info(sample_id)
    insert_person_info(person_info, doc)
    insert_qc(sample_id, sample_type, sample_path, doc)
    insert_result(json_data, doc)
    insert_drug(json_data, doc)
    insert_drug_instruction(json_data, sample_path, doc)
    insert_chem(sample_path, doc)
    save_doc(sample_id, person_info, doc, save_path)

def main():
    (options, args) = parser.parse_args()
    sample_id, sample_type, out_path,  = options.sample, options.sample_type, options.out_path,
    mutation_path = options.mutation_path
    reCmp = re.compile(r'\w+-\w')
    if not sample_id:
        if not args:
            print('请输入编号')
            exit()
        word = args[0].rstrip('/')
        if not re.match(reCmp, word):
            print('请输入正确编号')
            exit()
        sample_id = word.split('-')[0]
        sample_type = word.split('-')[1]
    sample_path = os.path.join('D:\work\\try\daan\S170002743-Tumor' )
    doc = Document(TEMPLE_FILE)
    person_info = get_person_info(sample_id)
    insert_person_info(person_info, doc)
    insert_qc(sample_id, sample_type, sample_path, doc)
    json_data = load_json('D:\work\\try\daan\\demo.json')
    insert_result(json_data, doc)
    insert_drug(json_data, doc)
    insert_drug_instruction(json_data, sample_path, doc)
    insert_chem(sample_id, DATA_PATH, doc)
    save_doc(sample_id, person_info, doc, save_path=None)

if __name__ == "__main__":
    json_data = load_json('D:\work\\try\daan\\demo.json')
    daan_run('S170002743', 'Tumor', json_data)

# doc = Document('D:\work\\try\daan\\temple_daan.docx')
# for pa in doc.paragraphs:
#     if pa.text == '接收日期：':
#         print (pa.text)

# person_table = doc.tables[3]
# for cell in person_table._cells:
#     if '姓名' in cell.text:
#         cell.paragraphs[0].add_run('aa')
# doc.save('D:\work\\try\daan\qqq.docx')

# for cell in doc.tables[3].row_cells(2):
#     print (cell.text)

# with open(JSON_FILE, encoding='utf8') as f:
#     json_data = json.load(f)
#
# for i in json_data:
#     print (i)

