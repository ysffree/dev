from __future__ import print_function
from collections import OrderedDict
from docx import Document
from sys import argv
from math import floor
from docx.shared import Inches
import docx
import sys
import os
import json
import traceback
from docx.oxml.shared import OxmlElement, qn
from optparse import OptionParser
from docx.enum.text import WD_ALIGN_PARAGRAPH


HEADLINES = (
    u"铂类",
    u"蒽环素类",
    u"芳香化酶抑制剂",
    u"鬼臼素类",
    u"抗雌激素类",
    u"嘧啶类似物",
    u"嘌呤类似物",
    u"烷化剂",
    u"喜树碱类",
    u"叶酸类似物",
    u"长春花生物碱类",
    u"紫杉烷类",
)

LOCAL_PATH = os.path.split(os.path.realpath(__file__))[0]
INSTRUCTION_UP = os.path.join(LOCAL_PATH, 'config', 'instruction_up.txt')
INSTRUCTION_DOWN = os.path.join(LOCAL_PATH, 'config', 'instruction_down.txt')
DEFAULT_FINE_FILE = os.path.join(LOCAL_PATH, 'config', 'fine.txt')
GENE_JSON = os.path.join(LOCAL_PATH, 'config', 'chem.genelist')
HEADERS_FILE = os.path.join(LOCAL_PATH, 'config', 'headers.json')

def get_headers(jsonfile):
    #print(os.path.abspath(HEADERS_FILE))
    headers = json.loads(open(HEADERS_FILE, encoding='utf-8').read(),object_pairs_hook=OrderedDict)
    res = json.loads(open(jsonfile, encoding='utf-8').read())
    # print('res', res)
    for header in headers:
        for title in header:
            if title not in res[0]:
                break
        else:
            break
    return header

def get_versions():
    dname = os.path.join(os.path.dirname(__file__), "version")
    print ("Version Lists:", file=sys.stderr)
    print (*os.listdir(dname), sep=", ", file=sys.stderr)


def get_chem_genelist_file(report_version="SmartOnco"):
    return os.path.join(os.path.dirname(__file__), "version", report_version, "chem.genelist")


def load_chem_genelist(genelist_file):
    genes = set()
    # genelist_file = get_chem_genelist_file(report_version)
    if not os.path.exists(genelist_file):
        print ("Warning: report_version:%s chem list file not exists" ,file=sys.stderr)
        return ()
    with open(genelist_file) as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            genes.add(line)
    return genes


def set_cell_vertical_alignment(cell, align="center"):
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')
        tcValign.set(qn('w:val'), align)
        tcPr.append(tcValign)
        return True
    except:
        traceback.print_exc()
        return False


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    if isinstance(text, str):
        return p.add_run(text)
#     r = set_cell_text_center(cell, text)
#     return set_cell_vertical_alignment(cell)
#
#
# def set_cell_text_center(cell, text, align="left"):
#     p = cell.paragraphs[0] # add_paragraph()
#     # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    # p.alignment = WD_ALIGN_PARAGRAPH.LEFT
#     if isinstance(text, str):
#         text = text.decode("utf-8")
#     return p.add_run(text)


def get_row(dic):
    rows = 0
    if isinstance(dic, dict) and dic:
        for name in dic:
            sub_dic = dic[name]
            rows += get_row(sub_dic)
    else:
        rows += 1
    return rows


def merge_cols(table, col, row_start, row_end):
    start_cell = table.cell(row_start, col)
    if row_start == row_end:
        return start_cell
    end_cell = table.cell(row_end, col)
    return start_cell.merge(end_cell)


def render_table(table, dic, start_row=1, col=0, headlines=None):
    cur_row = start_row
    keys = headlines if (headlines and col == 0) else sorted(dic.keys())
    for key in keys:
        if key not in dic:
            print (key, "not in list", file=sys.stderr)
            continue
        try:
            rows = get_row(dic[key])
            cell = merge_cols(table, col, cur_row, cur_row + rows - 1)
        except:
            print (keys, "::", key, "--", col, file=sys.stderr)
            print ("MERGE ROW", col, cur_row, cur_row + rows -1, file=sys.stderr)
            raise
        set_cell_text(cell, key)
        val = dic[key]
        if isinstance(val, dict):
            render_table(table, val, cur_row, col+1, headlines)
        cur_row += rows


def render(doc, table_info, headers):
    cols = len(headers.values())
    rows = get_row(table_info)
    # print ("total:", rows, file=sys.stderr)
    tbl = doc.add_table(rows+1, cols)
    # tbl = doc.add_table(rows+1, cols, style="chemicaldrug")
    tbl.style = 'chem table' # Normal Table'
    # tbl.style = 'Table Grid' # Normal Table'
    tbl.autofit = False

    # write header
    for i, key in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_text(cell, headers[key])
    render_table(tbl, table_info, headlines=HEADLINES)
    tbl.style = 'chem table' # Normal Table'


def get_table_info(jsonfile, headers, genelist=()):
    if not genelist or not jsonfile:
        return {}

    res = json.loads(open(jsonfile,encoding='utf-8').read())
    dic = {}
    gene = list(headers.keys())[list(headers.values()).index('检测基因')]
    for info in res:
        if info[gene] not in genelist:
            continue
        sub_dic = dic
        for header in headers:
            val = info.get(header, None)
            if not val or val == "-":
                break
            if val not in sub_dic:
                sub_dic[val] = {}
            sub_dic = sub_dic[val]
    return dic

def fine_file(filename = DEFAULT_FINE_FILE):
    return filename

def fine_words(filename):
    fine = set()
    with open(filename, encoding= 'utf-8') as f:
        for line in f:
            if not line or line.startswith('#'):
                continue
            line = line.strip('﻿')
            line = line.strip()
            fine.add(line)
    # print (fine)
    return fine

def is_strong(cell, fine_set):
    a = 0
    for word in cell.split('；'):
        if word not in fine_set:
            a += 1
    return a == 0

def red_table(table, fine_set):
    for row in table.rows:
        if is_strong(row.cells[-1].text, fine_set):
            for cell in row.cells[3:]:
                text = cell.text
                cell.text = ''
                cell.paragraphs[0].add_run(text, style='strong')

def strong_table(table, fine_set):
    poison_list = []
    effect_list = []
    list_all = []
    for row in table.rows:
        if is_strong(row.cells[-1].text, fine_set):
            if '药效增强' in row.cells[-1].text or '药物响应增强' in row.cells[-1].text:
                if row.cells[1].text not in effect_list:
                    effect_list.append(row.cells[1].text)
            if '毒副作用减弱' in row.cells[-1].text:
                if row.cells[1].text not in poison_list:
                    poison_list.append(row.cells[1].text)
    posion_str = '、'.join(poison_list)
    effect_str = '、'.join(effect_list)
    list_all.append(effect_str)
    list_all.append(posion_str)
    return list_all

def table_width(table, *width):
    if not width:
        print('width is empty')
        exit()
    else:
        cols = len(table.columns)
        _width = len(width)
        if cols >= _width:
            print (cols,_width)
            print (int(cols/_width))
            width = tuple(list(width) * floor(cols/_width) + list(width)[0:cols%_width])
        else:
            width = tuple(list(width)[0:cols])
        table.autofit = False
        for col in range(cols):
            table.cell(0, col)._tc.width = Inches(width[col]/2.45)

# def get_report_tempate(report_version="SmartLung"):
#     filepath = os.path.join(os.path.dirname(__file__), "version", report_version, "template.docx")
#     if not os.path.exists(filepath):
#          print ("can't find %s " % filepath, file=sys.stderr)
#          return None
#     return filepath


# def get_doc(template=None, report_version="SmartLung"):
#     if not template:
#         template = get_report_tempate(report_version)
#     if not template:
#        print ("not template found!", file=sys.stderr)
#        return docx.Document()
#     print ("find tempate: %s!" % template, file=sys.stderr)
#     return docx.Document(template)

def get_instruction(doc, file=None,style=None):
    with open(file, encoding= 'utf-8') as f:
        for line in f:
            if not line or line.startswith('#'):
                continue
            line = line.strip('﻿')
            line = line.strip()
            doc.add_paragraph(line,style)

def walk_dir(path, keyword):
    for (root, subdirs, files) in os.walk(path):
        for file in files:
            if keyword in file:
                return os.path.join(root, file)

# def get_chem_json(simple_id):
#     chem_file = walk_dir('/.../../', 'genotype.fix_germline.json')
#     chem_json = json.loads(open(chem_file, encoding='utf-8').read())
#     return chem_json


def insert_chem(sample_path, doc):
    # doc = docx.Document() if not template else docx.Document("template.docx")
    chem_json_file = walk_dir(sample_path, 'genotype.json')
    #chem_json_file = 'D:\\smaple\\S170001936-Tumor\\germline.json'
    # doc = get_doc(template=None, report_version=report_version)
    doc.add_paragraph(u'2、化疗药物相关基因检测结果', "p5")
    doc.add_paragraph(u'本检测涉及19种常见化疗药物的毒性和药效评估内容，检测结果解读来源于PharmGKB 数据库，结果如下表：', 'p1')
    get_instruction(doc, INSTRUCTION_UP,'p6')
    # doc.add_paragraph(u'可能会出现一个药物与多个位点有关，药物具体的疗效和毒副作用需要综合判断','p16')
    genelist = load_chem_genelist(GENE_JSON)
    HEADERS = get_headers(jsonfile=chem_json_file)
    table_info = get_table_info(jsonfile=chem_json_file, headers=HEADERS, genelist=genelist)
    # if not table_info:
    #     print ("Warning: no chem drug get by current report_version:%s" % report_version, file=sys.stderr)
    #     return
    doc.add_paragraph()
    doc.add_paragraph('所有药物检测结果','p3')
    render(doc, table_info, HEADERS)
    table = doc.tables[-1]
    table_width(table, 2.24, 2.17, 2.0, 2.44, 2.34, 1.9, 2.92)
    filename = fine_file()
    fine_set = fine_words(filename)
    list_all = strong_table(table, fine_set)
    red_table(table, fine_set)
    # pa = doc.paragraphs[3]
    # pa.insert_paragraph_before('\u258E 药效/药物响应增强药物','drug name')
    # drug1 = pa.insert_paragraph_before('{}'.format(list_all[0]))
    # drug1.style='p1'
    # pa.insert_paragraph_before('\u258E 毒副作用减弱药物','drug name')
    # drug2 = pa.insert_paragraph_before('{}'.format(list_all[1]))
    # drug2.style='p1'
    doc.add_paragraph()
    get_instruction(doc, INSTRUCTION_DOWN,'p6')
    doc.add_page_break()
    # outname = os.path.basename(chem_json_file).split(".")[0]
    # doc.save("%s.chem-%s.docx" % (outname, report_version))


# if __name__ == "__main__":
    # report_version = "SmartOnco"
    # args = sys.argv[1:]
    # if args and args[0].startswith("-"):
    #     report_version = args[0][1:]
    #     args = args[1:]
    # if not args:
    #     print ("Usage: %s [-version] jsonfile" % __file__, file=sys.stderr)
    #     get_versions()
    #     exit(1)
    # for arg in args:
    #     run(arg, report_version=report_version)
