#!/usr/bin/env python3
# coding:utf-8

from docx import Document
from ddpcr_info import get_ddpcr_info
from person_info import get_person_info
from sys import argv
from docx.shared import Inches
from lxml import etree
import zipfile
import tempfile
import os
import shutil
import time
import xlrd
from collections import OrderedDict

# durg_info_path = os.path.join(os.path.split(os.path.realpath(__file__))[0], '/Main.xls')
# durg_info_path = 'C:\\Users\ysf\Desktop\data\检测\\main.xls'

# def durg_name(name_str):
#     name_list = name_str.split('|')
#     if name_list[1] != '【待定】':
#         name = name_list[1]
#     else:
#         name = name_list[0]
#     return name
#
# def drug_info(mutation):
#     date = xlrd.open_workbook(durg_info_path)
#     table = date.sheets()[0]
#     if mutation == '19del':
#         mutation = 'exon19del'
#     else:
#         mutation = 'p.' + mutation
#     info = OrderedDict()
#     for row in range(table.nrows):
#         if table.cell(row, 21).value == mutation or table.cell(row, 22).value == mutation:
#             name = durg_name(table.cell(row, 8).value)
#             info[name] = table.cell(row, 26).value
#     return info
#
# def judge_result(ddpcr_info):
#     info = OrderedDict()
#     mutation_result = OrderedDict()
#     for sid in ddpcr_info:
#         mutation_result[sid["突变"]] = sid["定性结果"]
#     if 'T790M' in mutation_result and mutation_result['T790M'] == '阳性':
#         info = drug_info('T790M')
#     else:
#         for mutation in  mutation_result:
#             if mutation_result[mutation] == '阳性':
#                 info = drug_info(mutation)
#     return info
#
#
# def drug_table(doc, ddpcr_info):
#     drug_info = judge_result(ddpcr_info)
#     tbl = doc.add_table(1,2,'mutation result')
#     tbl.cell(0, 0).text = '药物'
#     tbl.cell(0, 1).text = '提示'
#     row = 0
#     if not drug_info:
#         tbl.add_row()
#         tbl.cell(1, 0).text = '没有用药提示'
#     else:
#         for drug in drug_info:
#             tbl.add_row()
#             row += 1
#             tbl.cell(row,0).text = drug
#             tbl.cell(row,1).text = drug_info[drug]

def mutation_handle(ddpcr_info):
    mutation_result = OrderedDict()
    for sid in ddpcr_info:
        mutation_result[sid['突变']] = sid["定性结果"]
    return mutation_result

def drug_table(doc, ddpcr_info):
    mutation_result = mutation_handle(ddpcr_info)
    mutation_list = list(mutation_result.keys())
    #print(mutation_list)
    mutation_list.remove('T790M')
    positive_list = []
    # mutation_str = '\n'.join(mutation_list)
    for mutation in mutation_list:
         if mutation_result[mutation] == '阳性':
                positive_list.append(mutation)
    mutation_str = '\n'.join(positive_list)
    result_list = mutation_result.values()
    if '阳性' not in result_list:
        tbl = doc.add_table(1, 1, 'mutation result')
        tbl.cell(0,0).text = '本次检测未发现用药提示'
    else:
        tbl = doc.add_table(1, 3, 'mutation result')
        tbl.cell(0, 0).text = 'EGFR体细胞突变'
        tbl.cell(0, 1).text = '敏感'
        tbl.cell(0, 2).text = '耐药'
        i = 1
        for mutation in mutation_list:
            if mutation_result[mutation] == '阳性':
                tbl.add_row()
                tbl.cell(1, 0).text = mutation_str
                tbl.cell(1, 1).paragraphs[0].add_run('吉非替尼\n阿法替尼\n厄洛替尼\n埃克替尼\n奥希替尼', 'red')
                tbl.cell(1, 2).text = '--'
                i = 2
                break
        if mutation_result['T790M'] == '阳性':
            tbl.add_row()
            tbl.cell(i, 0).text = 'T790M'
            tbl.cell(i, 1).paragraphs[0].add_run('奥希替尼', 'red')
            tbl.cell(i, 2).paragraphs[0].add_run('吉非替尼\n阿法替尼\n厄洛替尼\n埃克替尼', 'blue')


def extract_zip_to_tmpdir(zfile):
    tmp_dir = tempfile.mkdtemp()
    zfile.extractall(tmp_dir)
    return tmp_dir


def modify_header(tmp_dir, sample_id,  report_time):
    header_path = os.path.join(tmp_dir, "word/header1.xml")
    xml = etree.parse(header_path)
    root = xml.getroot()
    for tag in root.iter():
        if str(tag.text).find('样本编号') != -1:
            tag.text += '：' + sample_id
        elif str(tag.text).find('报告日期') != -1:
            tag.text += '：' + report_time
    xml.write(header_path)


def person_info_table(doc, person_info):
    tbl = doc.add_table(rows=3,cols=4,style='baseinfo')
    for i, key in enumerate(person_info):
        if i < 3:
            tbl.cell(i, 0).text = key
            tbl.cell(i, 1).text = person_info[key]
        else:
            tbl.cell(i-3, 2).text = key
            tbl.cell(i-3, 3).text = person_info[key]
    doc.add_paragraph()

def general_result_table(doc, ddpcr_info):
    row = len(ddpcr_info)
    tbl = doc.add_table(row+1, 2, 'mutation result')
    tbl.cell(0, 0).text = '检测项目'
    tbl.cell(0, 1).text = '结果'
    for i, sid in enumerate(ddpcr_info):
        tbl.cell(i+1, 0).text = 'EGFR ' + sid['突变']
        if sid['定性结果'] == '阴性':
            tbl.cell(i+1, 1).paragraphs[0].add_run(sid['定性结果'],'blue')
        else:
            tbl.cell(i + 1, 1).paragraphs[0].add_run(sid['定性结果'], 'red')
    doc.add_paragraph()


def add_pic(doc, sid):
    doc.add_paragraph('\u258E实验数据图','p6')
    sid["图片"].save("{0}.png".format(sid["突变"]))
    doc.add_picture("{0}.png".format(sid["突变"]), width=Inches(6.0), height=Inches(2.0))
    os.remove("{0}.png".format(sid["突变"]))

def detail_result_info(doc, ddpcr_info):
    for sid in ddpcr_info:
        small_title = '\u258EEGFR ' + sid['突变'] + '位点'
        doc.add_paragraph(small_title, 'drug name' )
        tbl = doc.add_table(2,5,'ddpcr result')
        # remove = ["突变","图片"]
        for i,title in enumerate(sid):
            if i < 5:
                tbl.cell(0,i).text = title
                tbl.cell(1,i).text = str(sid[title])
        if tbl.cell(1,4).text == '阳性':
            tbl.cell(1,4).text = ''
            tbl.cell(1,4).paragraphs[0].add_run(sid["定性结果"], 'red').font.bold = True
        doc.add_paragraph()
        add_pic(doc, sid)


def word_name(ddpcr_info,person_info,sample_id):
    num = len(ddpcr_info)
    if num == 1:
        name = "{0}-{1}-ddpcr-{2}.docx".format(sample_id, person_info["姓  名"], ddpcr_info[0]["突变"])
    else:
        name = "{0}-{1}-ddpcr-{2}位点.docx".format(sample_id, person_info["姓  名"], num)
    return name

def handle_id(simple_id):
    if '-'in simple_id:
        simple_id = simple_id.split('-')[0]
    return simple_id

def main():
    local_path = os.path.split(os.path.realpath(__file__))[0]
    # print (local_path)
    sample_id = argv[1]
    process_id = handle_id(sample_id)
    doc = Document(os.path.join(local_path,'temple_v2.docx'))
    doc.add_paragraph('\u2589 基本信息', "title")
    try:
        person_info = get_person_info(process_id)
        print ('aa')
    except Exception as e:
       # doc.add_paragraph('\u2589 基本信息', "title")
        person_info = dict.fromkeys(["姓  名", "性  别", "年  龄", "病理诊断", "样本类型", "送检项目"], '')
        print(e)
    person_info_table(doc, person_info)
    doc.add_paragraph('\u2589 本次检测概览',"title")
    ddpcr_info = get_ddpcr_info(sample_id)
    doc.add_paragraph('本检测是在QX200平台上进行的数字PCR检测，针对下表中EGFR基因{}个位点所设计。受技术手段所限，本实验'
                      '结果不能判断在该位点以外的区域（同一基因的其他区域）的突变情况。'.format(len(ddpcr_info)), 'p1')

    doc.add_paragraph('\u258E实验检测结果','drug name')
    general_result_table(doc, ddpcr_info)
    doc.add_paragraph('\u258E用药提示', 'drug name')
    drug_table(doc, ddpcr_info)
    doc.add_page_break()
    doc.add_paragraph('\u2589 详细实验结果', "title")
    detail_result_info(doc, ddpcr_info)
    if os.path.exists(os.path.join(local_path, "report")) == True:
        pass
    else:
        os.mkdir(os.path.join(local_path, "report"))
    name = word_name(ddpcr_info,person_info,process_id)
    doc.save(os.path.join(local_path, "report/{}".format(name)))
    print (os.path.join(local_path, "report/{}".format(name)))
    docxfile = os.path.join(local_path,"report/{}".format(name))
    zfile = zipfile.ZipFile(docxfile, "r")
    tmp_dir = extract_zip_to_tmpdir(zfile)
    modify_header(tmp_dir, process_id, time.strftime("%Y-%m-%d"))
    namelist = zfile.namelist()
    with zipfile.ZipFile(os.path.join(local_path, "report/{}".format(name)),"w") as out:
        for filename in namelist:
            out.write(os.path.join(tmp_dir, filename), filename)
    shutil.rmtree(tmp_dir)

if __name__ == "__main__":
    main()
