from docx import Document
temple = Document('D:\work\\try\wordtry\smartonco156.docx')

# # 段落标题
# temple.add_paragraph('\ue136 基本信息', "title")
#
# # 基本信息表格
# temple.add_table(rows=4,cols=4,style='baseinfo')
# temple.add_paragraph()
# temple.add_paragraph()
# temple.add_paragraph()
#
# # 段落标题
# temple.add_paragraph('\ue133 用药检测结果', "title")
#
#
# # 正文 首行缩进 宋体/Cambria 10.5 黑色 不加粗 段后距10pt 段间距 1.15
# temple.add_paragraph('本次检测包含156个靶向药物相关基因，共发现10个当前癌种的用药提示，0个其他癌种的用药提示，见下表：', 'p1')
# temple.add_paragraph()
#
# # 癌种用药提示表格
# title = temple.add_table(rows=1,cols=7,style='table title')
# title.cell(0,0).merge(title.cell(0,6)).text = '当前癌种用药提示'
# temple.add_table(rows=7,cols=7,style='target or use')
# temple.add_paragraph()
#
# # 正文 首行不缩进 宋体/Cambria 10 黑色 加粗 段后间距10pt 段间距 1.15
# temple.add_paragraph('注：', 'p2')
# temple.add_paragraph('1. 敏感/耐药：来源于FDA药品说明书(Drug Label) 或肿瘤临床指南。', 'p2')
# temple.add_paragraph('2. 可能敏感/可能耐药：来源于文献或临床试验报道。', 'p2')
# temple.add_paragraph('3. 本报告暂时只对单个位点进行药物敏感性判断，然而具体药物的敏感性可能受多个基因多个位点的影响，最终结果需要临床医生根据实际情况综合判定。', 'p2')
#
# # 正文 首行不缩进 宋体/Cambria 11 黑色 不加粗 段后距10pt 段间距 1.15
# temple.add_paragraph('本次检测共覆盖114种靶向药物，见下表：', 'p3')
# temple.add_paragraph()
#
# # CFDA或FDA批准的药物表格
# temple.add_table(rows=4, cols=4, style='druginfo1')
# # 未上市的药物表格
# temple.add_table(rows=4, cols=4, style='druginfo2')
#
# temple.add_page_break()
#
# # 段落标题
# temple.add_paragraph('\ue8c0 靶向药物基因检测结果总览', "title")
#
# # 靶向药物基因检测结果总览表格
# temple.add_table(rows=7,cols=7,style='target or use')
# temple.add_page_break()
#
# # 段落标题
# temple.add_paragraph('\ue8c0 基因检测结果临床意义', "title")
# # temple.add_paragraph()
#
#
# # 正文 首行不缩进 宋体/Cambria 10.5 蓝色 加粗 段后距10pt 段间距 1.15
# temple.add_paragraph('检测基因：EGFR', "p4")
#
# # 正文 首行不缩进 段后间距10pt 段间距 1.15
# p = temple.add_paragraph(style="p7")
# # 正文 宋体/Cambria 10.5 蓝色 加粗
# p.add_run('基因说明：',"r3")
# # 正文 宋体/Cambria 10 黑色 不加粗
# p.add_run('EGFR（表皮生长因子受体）是原癌基因cerbB1的表达产物，是表皮生长因子受体（HER）家族成员之一，普遍表达于人体的表皮细胞和基质细胞',"r1")
#
#
#
# # 基因检测结果临床意义，单个基因表
# temple.add_table(rows=2,cols=7,style='single gene')
# temple.add_paragraph()
#
#
# k = temple.add_paragraph(style="p7")
# # 正文 宋体/Cambria 四号 红色 加粗
# k.add_run('肺癌',"r2")
# k.add_run('中临床意义描述',"r3")
#
# # 正文 首行缩进 宋体/Cambria 10 黑色 不加粗 段后距10pt 段间距 1.15
# temple.add_paragraph('p.T790M点突变发生于EGFR第20外显子。 EGFR基因外显子20上的插入突变（除p.A763_Y764insFQEA外）以及T790M突变代', "p8")
# temple.add_paragraph()
# temple.add_paragraph('-'*125 ,'p5')
# temple.add_page_break()
#
# # 段落标题
# temple.add_paragraph('\ue8c0 化疗药物相关基因检测结果', "title")
#
# temple.add_paragraph('本检测涉及20种常见化疗药物的毒性和药效评估内容，检测结果解读来源于PharmGKB 数据库，结果如下表：', 'p1')
#
# # 化疗药物相关基因检测结果表格
# chemical= temple.add_table(rows=18,cols=7,style='chemical drug')
# temple.add_paragraph()
# temple.add_paragraph()
#
# # 正文 首行不缩进 宋体/Cambria 10 蓝色 加粗 段后距10pt 段间距 1.15
# temple.add_paragraph('说明：', "p5")
# # 正文 首行不缩进 宋体/Cambria 10 黑色 不加粗 段后间距10pt 段间距 1.15
# temple.add_paragraph('•	基因名称均采用 NCBI-Gene 官方命名(Official Symbol)', "p6")
# temple.add_paragraph('•	证据等级(Level) 的划分：依据 PharmGKB 网站\nhttp://www.pharmgkb.org/page/clinAnnLevels', "p6")
# temple.add_paragraph('–	1A：注释基于被医学会认可的指南或经某些重大卫生系统认可的结论', "p6")
# temple.add_page_break()
#
# # 段落标题
# temple.add_paragraph('\ue8c0 体细胞突变检测结果', "title")
#
# # 体细胞突变检测结果表格（第二行加粗待解决）
# temple.add_table(rows=7,cols=5,style='mutation result')
# temple.add_page_break()
#
# # 段落标题
# temple.add_paragraph('\ue8c0 质控统计', "title")
#
# # 质控统计表格
# temple.add_table(rows=7,cols=2,style='qc')
#
# # 页眉
# temple.add_table(rows=2,cols=2,style='head')
#
# # 页脚
# temple.add_paragraph('SmartOnco 肿瘤无创基因检测报告-1', "foot")
#
# # EGFR/EAR 基本信息表格
# temple.add_table(rows=2,cols=4,style='baseinfo2')
#
# # 正文 首行不缩进 宋体/Cambria 12 黑色 加粗 段后距10pt 段间距 1.15
# temple.add_paragraph('实验检测结果：', "p9")
#
# # 实验检测结果表
# temple.add_table(rows=5,cols=2,style='testing result')
#
# # 正文 首行不缩进 宋体/Cambria 12 黑色 不加粗 段后距10pt 段间距 1.15
# temple.add_paragraph('•	本检测仅包含EGFR基因（NM_005228.3）18-21外显子内29种已知常见突变，并不能排除该样本带有其他位点的突变；', "p10")
# temple.add_paragraph('•	突变DNA最低检测限约为总DNA样本含量1%以上；', "p10")



# # p11：正文 居中 宋体/Cambria 14 黑色 加粗 段前距10pt 段后距10pt 段间距 1.5
# temple.add_paragraph('循环肿瘤DNA（ctDNA）高通量基因检测报告', "p11")
# # baseinfo3：佰美基本信息表格
# temple.add_table(rows=6,cols=3,style='baseinfo3')
# # p12：正文 首行不缩进 宋体/Cambria 14 黑色 加粗 段前距10pt 段后距10pt 段间距 1.5
# temple.add_paragraph('一、 检测项目说明', "p12")
# # p13：正文 首行缩进 宋体/Cambria 12 黑色 加粗 段后距5pt 段间距 1.5 底纹黄色
# temple.add_paragraph('□ 肿瘤驱动基因检测', "p13")
# # genetest：佰美其他表格
# temple.add_table(rows=6,cols=4,style='genetest')
#
# # p14：正文 首行缩进 宋体/Cambria 12 红色 不加粗 段前距5pt 段间距 1.5
# temple.add_paragraph('本次检测在Illumina Hiseq X Ten平台上进行的高通量测序，覆盖上述基因的热点突变、插入缺失、基因融合、基因扩增等变异类型，平', "p14")
#
# temple.add_paragraph('二、	肿瘤驱动基因检测结果与说明', "p12")
#
# # p9：正文 首行不缩进 宋体/Cambria 12 黑色 加粗 段后距10pt 段间距 1.15
# temple.add_paragraph('1、	所检出突变具体信息', "p9")
#
# # p15：正文 首行缩进 宋体/Cambria 10.5 黑色 不加粗 段后距10pt 段间距 1.5
# temple.add_paragraph('本次检测包含136个靶向药物相关基因，共发现5个用药提示，剩余135个基因未检测到用药提示相关变异。突变信息如下：', "p15")
#
# temple.add_table(rows=6,cols=8,style='genetest')
#
# # p17：正文 首行不缩进 宋体/Cambria 9 黑色 不加粗 段前距10pt 段后距10pt 段间距 1.5
# temple.add_paragraph('[1] FDA，证据来源于FDA药品使用说明书。[2] NCCN，证据来源于NCCN指南。[3] 临床（前）研究：尚处于临床研究阶段，有文献报告该基因突变可能与药物疗效有关。', "p17")
#
# temple.add_paragraph('2、	检测结果说明', "p9")
#
# # p16：正文 首行缩进 宋体/Cambria 10.5 黑色 加粗  段后距10pt 段间距 1.5
# temple.add_paragraph('1)  基因:ARID1A  变异形式:ARID1A基因第441位编码子丙氨酸变为缬氨酸', "p16")
# temple.add_paragraph('    基因说明：', "p16")
# temple.add_paragraph('    原癌基因HER2编码的185kDa的细胞膜受体，为表皮生长因子受体家族成员之一。该家族包括ERBB1', "p15")
#
# temple.add_paragraph('3、	参考文献：', "p9")
# temple.add_paragraph('Druglabel:ado-trastuzumab emtansine_Revised:04/2016', "p15")

temple.add_table(rows=1, cols=4, style='druginfo1')
temple.add_table(rows=4, cols=4, style='druginfo3')
temple.add_table(rows=1, cols=4, style='druginfo1')
temple.add_table(rows=4, cols=4, style='druginfo4')


temple.save('D:\work\\try\wordtry\sample.docx')



# paragraph
# title：标题 首行不缩进 黑体/Cambria 三号 黑色 加粗 段后距10pt 段间距 1.15
# p1：正文 首行缩进 宋体/Cambria 10.5 黑色 不加粗 段后距10pt 段间距 1.15
# p2：正文 首行不缩进 宋体/Cambria 10 黑色 加粗 段后间距10pt 段间距 1.15
# p3：正文 首行不缩进 宋体/Cambria 11 黑色 不加粗 段后距10pt 段间距 1.15
# p4：正文 首行不缩进 宋体/Cambria 10.5 蓝色 加粗 段后距10pt 段间距 1.15
# p5：正文 首行不缩进 宋体/Cambria 10 蓝色 加粗 段后距10pt 段间距 1.15
# p6：正文 首行不缩进 宋体/Cambria 10 黑色 不加粗 段后间距10pt 段间距 1.15
# p7：正文 首行不缩进 段后间距10pt 段间距 1.15
# p8：正文 首行缩进 宋体/Cambria 10 黑色 不加粗 段后距10pt 段间距 1.15
# p9：正文 首行不缩进 宋体/Cambria 12 黑色 加粗 段后距10pt 段间距 1.15
# p10：正文 首行不缩进 宋体/Cambria 12 黑色 不加粗 段后距10pt 段间距 1.15
# p11：正文 居中 宋体/Cambria 14 黑色 加粗 段前距10pt 段后距10pt 段间距 1.5
# p12：正文 首行不缩进 宋体/Cambria 14 黑色 加粗 段前距10pt 段后距10pt 段间距 1.5
# p13：正文 首行缩进 宋体/Cambria 12 黑色 加粗 段后距10pt 段间距 1.5 底纹黄色
# p14：正文 首行缩进 宋体/Cambria 12 红色 不加粗 段前距10pt 段间距 1.5
# p15：正文 首行缩进 宋体/Cambria 10.5 黑色 不加粗 段后距10pt 段间距 1.5
# p16：正文 首行缩进 宋体/Cambria 10.5 黑色 加粗  段后距10pt 段间距 1.5
# p17：正文 首行不缩进 宋体/Cambria 9 黑色 不加粗 段前距10pt 段后距10pt 段间距 1.5
# foot: 页脚 宋体/Cambria 9 居中 蓝色 不加粗 段后距10pt 段间距 1.15

# text/run
# r1：正文 宋体/Cambria 10 黑色 不加粗
# r2：正文 宋体/Cambria 四号 红色 加粗
# r3：正文 宋体/Cambria 10.5 蓝色 加粗
# r4：正文 宋体/Cambria 10.5 黑色 加粗
# r5：正文 宋体/Cambria 10.5 黑色 不加粗
# strong:正文 宋体/Cambria 10 红色 加粗

# table：
# baseinfo：基本信息表格
# table title：表格头（跟癌症用药提示表格配合使用）
# target or use：癌症用药提示表格/靶向药物基因检测结果总览表
# druginfo1：CFDA或FDA批准的药物表格
# druginfo2：未上市的药物表格
# single gene：基因检测结果临床意义中单个基因表
# chemical drug：化疗药物相关基因检测结果表格
# mutation result：体细胞突变检测结果表格（第二行加粗待解决）
# qc：质控统计表格
# head：页眉表格
# baseinfo2：EGFR/EAR 基本信息表格
# testing result：实验检测结果表
# baseinfo3：佰美基本信息表格
# genetest：佰美其他表格
# druginfo3：CFDA或FDA批准的药物表格中间部分
# druginfo4：未上市的药物表格中间部分


# temple.add_page_break()