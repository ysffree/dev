# paragraph
# title：宋体 四号 居中 加粗
# p1：宋体 四号 左对齐 加粗
# p2：宋体小四 首行缩进 加粗
# p3：宋体小四 首行缩进 不加粗 字体红色
# p4：宋体小四 首行不缩进 加粗
# p5：宋体五号 首行缩进 不加粗
# p6：宋体小五 首行不缩进 不加粗
# p7：宋体五号 首行缩进  加粗
# p8：宋体五号 每行缩进  不加粗
# p9：宋体五号 首行不缩进  加粗
# p10：宋体 Wingdings 五号 首行缩进 不加粗 段间距0
# p11：宋体五号 首行不缩进 每行缩进  不加粗
# p12：宋体五号 首行缩进 不加粗 段间距0

# run
# r1：宋体小四 加粗 底纹黄色

# table
# baseinfo：基本信息table
# other table：其他table

from docx import Document
temple = Document('D:\work\\try\\baimei\\temple_baimei.docx')
temple.add_paragraph('循环肿瘤DNA（ctDNA）高通量基因检测报告', 'title')
temple.add_table(5, 3, 'baseinfo')
temple.add_paragraph('一、	检测项目说明', 'p1')
pa = temple.add_paragraph('□ ', 'p2')
pa.add_run('肿瘤驱动基因检测', 'r1')
temple.add_table(10, 4, 'other table')
temple.add_paragraph('本次检测在Illumina Hiseq X Ten平台上进行的高通量测序，覆盖上述基因的热点突变、插入缺失、基因融合、', 'p3')
temple.add_paragraph('1、	所检出突变具体信息', 'p4')
temple.add_paragraph('本次检测包含50个靶向药物相关基因，共发现0个用药提示，剩余50个基因未检测到用药提示相关变异。突变信息如下：', 'p5')
temple.add_table(10, 4, 'other table')
temple.add_paragraph('[1] FDA，证据来源于FDA药品使用说明书。[2] NCCN，证据来源于NCCN指南。[3] 临床（前）研究：尚处于临床研究阶段，有文献报告该基因突变可能与药物疗效有关。', 'p6')
temple.add_paragraph('1)  基因: TP53  变异形式: TP53基因第234位编码子酪氨酸变为半胱氨酸', 'p7')
temple.add_paragraph('    基因说明：', 'p7')
temple.add_paragraph('    TP53基因是一种抑癌基因，其编码的蛋白P53是细胞生长、增殖和损伤修复的重要调节因子。在DNA损伤时，p53可使细胞G1/S', 'p8')
temple.add_paragraph('可能会出现一个药物与多个位点有关，药物具体的疗效和毒副作用需要综合判断', 'p9')
temple.add_paragraph('基因名称均采用 NCBI-Gene 官方命名(Official Symbol)', 'p10')
temple.add_paragraph('检测位点(rs 号)：NCBI 里对所有提交的 snp 进行分类考证之后都会给出一个 rs 号(也可称作参考 snp) 并给出 snp 的具体信息，包括前后序列、位置信息以及分布频率。', 'p10')
temple.add_paragraph('–	1A：注释基于被医学会认可的指南或经某些重大卫生系统认可的结论', 'p10')
temple.add_paragraph('–	1A：注释基于被医学会认可的指南或经某些重大卫生系统认可的结论', 'p10')
temple.add_paragraph('–	1A：注释基于被医学会认可的指南或经某些重大卫生系统认可的结论', 'p10')
temple.add_paragraph('–	1A：注释基于被医学会认可的指南或经某些重大卫生系统认可的结论', 'p10')
temple.add_paragraph('附录1：药物简介', 'p1')
temple.add_paragraph('阿法替尼：商品名：Gilotrif。2013年获得美国FDA批准上市。适用于携带有EGFR突变的转移性非小细胞肺癌(NSCLC)的一线治疗及HER2突变阳性的晚期乳腺癌患者。', 'p11')

temple.save('D:\work\\try\\baimei\\sample.docx')